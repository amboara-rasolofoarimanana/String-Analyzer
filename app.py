
# Importation des bibliothèques 
import streamlit as st
import pandas as pd
import altair as alt
import altair_saver
import numpy as np
import vl_convert as vlc

import re
import os
from io import BytesIO
import tempfile

from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table
from reportlab.lib.styles import getSampleStyleSheet


# Fonctions auxiliaires 

def trier_strings_par_numero(liste_strings):
    """
    Trie une liste de noms de strings en fonction de leur numéro de façon croissante.
    Les noms contenant des chiffres seront triés numériquement.
    Ceux ne contenant pas de chiffre seront mis à la fin.

    Args:
        liste_strings (list): Liste de noms de strings (chaînes de caractères).

    Returns:
        list: Liste triée.
    """
    def extract_num(s):
        match = re.search(r"\d+", s)
        return int(match.group()) if match else float('inf')  # 'total' en dernier
    return sorted(liste_strings, key=extract_num)

@st.cache_data
def lire_fichier(fichier):
    """
    Lit un fichier CSV ou Excel et retourne un dataframe.

    Args:
        fichier (UploadedFile ou str): Fichier à lire.

    Returns:
        pd.DataFrame: Données lues.

    Raises:
        ValueError: Si le format du fichier n’est pas pris en charge.
    """
    nom = fichier.name if hasattr(fichier, 'name') else str(fichier)
    extension = os.path.splitext(nom)[1].lower()

    if extension in ['.xls', '.xlsx']:
        return pd.read_excel(fichier)
    elif extension == '.csv':
        return pd.read_csv(fichier, encoding='utf-8')
    else:
        raise ValueError(f"Format de fichier non pris en charge : {extension}")

@st.cache_data
def traiter_fichier_onduleur(file):
    """
    Traite un fichier de production onduleur :
    - Renomme les colonnes
    - Convertit les types
    - Trie les lignes par date

    Args:
        file: Fichier CSV/Excel contenant les données de production.

    Returns:
        pd.DataFrame: Données formatées.
    """
    df = lire_fichier(file)
    nb_colonnes = len(df.columns)
    nb_strings = nb_colonnes - 2
    df.columns = ["time"] + [f"string {i}" for i in range(1, nb_strings + 1)] + ["total"]
    
    colonnes_a_convertir = [col for col in df.columns if col != "time"]
    df[colonnes_a_convertir] = df[colonnes_a_convertir].apply(pd.to_numeric, errors="coerce")
    df["time"] = pd.to_datetime(df["time"], errors="coerce")
    df = df.sort_values("time").reset_index(drop=True)
    return df

@st.cache_data
def traiter_fichier_carac(file):
    """
    Traite un fichier de caractéristiques strings :
    - Renomme les colonnes
    - Convertit les types numériques

    Args:
        file: Fichier CSV/Excel contenant les caractéristiques.

    Returns:
        pd.DataFrame: Données formatées.
    """
    df = lire_fichier(file)
    df.columns = ["string", "puissance unitaire", "nombre pv"]
    df["string"] = pd.to_numeric(df["string"], errors="coerce", downcast="integer")
    df["puissance unitaire"] = pd.to_numeric(df["puissance unitaire"], errors="coerce")
    df["nombre pv"] = pd.to_numeric(df["nombre pv"], errors="coerce", downcast="integer")
    return df

@st.cache_data
def traiter_fichier_irradiance(file):
    """
    Traite un fichier d'irradiance :
    - Renomme les colonnes
    - Convertit les types
    - Trie les lignes par date

    Args:
        file: Fichier CSV/Excel d'irradiance.

    Returns:
        pd.DataFrame: Données formatées.
    """
    df = lire_fichier(file)
    df.columns = ["time", "irradiance"]
    df["irradiance"] = pd.to_numeric(df["irradiance"], errors="coerce")
    df["time"] = pd.to_datetime(df["time"], errors="coerce")
    df = df.sort_values("time").reset_index(drop=True)
    return df

def sauvegarder_chart_png(chart, nom_fichier_png):
    """
    Sauvegarde un graphique Altair au format PNG à partir de sa spécification Vega-Lite.

    Args:
        chart (alt.Chart): Le graphique Altair.
        nom_fichier_png (str): Nom du fichier de sortie.

    Returns:
        str: Chemin vers le fichier PNG créé.
    """
    spec = chart.to_dict()
    png_data = vlc.vegalite_to_png(spec)
    with open(nom_fichier_png, "wb") as f:
        f.write(png_data)
    return nom_fichier_png

def generer_word(site, onduleur, date_debut, date_fin, img_barres, top3_df, flop3_df, df_alertes, img_evolution,
                 inclure_ratio=True, inclure_classement=True, inclure_analyse_suspect=True, inclure_evolution=True,logo_path=None):
    """
    Génère un rapport Word contenant les résultats d’analyse 

    Args:
        site (str): Nom du site.
        onduleur (str): Identifiant de l’onduleur.
        date_debut (str): Date de début de période.
        date_fin (str): Date de fin de période.
        img_barres (str): Chemin vers l’image du graphique des ratios.
        top3_df (pd.DataFrame): Top 3 des strings plus performants.
        flop3_df (pd.DataFrame): Bottom 3 des strings moins performants.
        df_alertes (pd.DataFrame): Données des strings suspects.
        img_evolution (str): Chemin vers l’image d’évolution.
        inclure_ratio (bool): Ajouter section ratio.
        inclure_classement (bool): Ajouter section classement.
        inclure_analyse_suspect (bool): Ajouter section alertes.
        inclure_evolution (bool): Ajouter graphique d’évolution.
        logo_path (str): Chemin vers le logo à insérer (optionnel).

    Returns:
        str: Chemin vers le fichier Word généré.
    """
    doc = Document()
    
   
    def add_text_paragraph(text):
        p = doc.add_paragraph(text)
        run = p.runs[0]
        run.font.name = 'Calibri'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')  
        run.font.color.rgb = RGBColor(0, 0, 0)

    
    def add_heading2(text):
        p = doc.add_heading(text, level=1)
        run = p.runs[0]
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 51, 102)  # bleu foncé (#003366)

    def add_centered_image(image_path, width_in_inches=6.5):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width_in_inches))

    # === Ajout du logo ===
    if logo_path:
        paragraph_logo = doc.add_paragraph()
        paragraph_logo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
        run_logo = paragraph_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1))  # adapte la taille selon ton logo
        
    
     # === Titre principal ===
    titre = doc.add_heading(f"Rapport de performance {site}", 0)
    titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = titre.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # noir
    run.font.bold = True

    # === Infos principales ===
    add_text_paragraph(f"Onduleur : {onduleur}")
    add_text_paragraph(f"Période : {date_debut} à {date_fin}")



     # === Ratio par string ===
    if inclure_ratio:
        add_heading2("Ratio kWh/kWc par string")
        add_centered_image(img_barres)


     # === Classement === 
    if inclure_classement:
        add_heading2("Classement des strings")

        add_text_paragraph("Top 3 - Strings plus performants")
        table_top = doc.add_table(rows=1, cols=2)
        table_top.style = 'Light Grid'
        table_top.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_top.autofit = True
        hdr_cells = table_top.rows[0].cells
        hdr_cells[0].text = 'String'
        hdr_cells[1].text = 'Ratio kWh/kWc'
        for _, row in top3_df.iterrows():
            row_cells = table_top.add_row().cells
            row_cells[0].text = str(row["string_label"])
            row_cells[1].text = f"{row['ratio kWh/kWc']:.2f}"

        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()



        add_text_paragraph("Bottom 3 - Strings moins performants")
        table_flop = doc.add_table(rows=1, cols=2)
        table_flop.style = 'Light Grid'
        hdr_cells = table_flop.rows[0].cells
        hdr_cells[0].text = 'String'
        hdr_cells[1].text = 'Ratio kWh/kWc'
        for _, row in flop3_df.iterrows():
            row_cells = table_flop.add_row().cells
            row_cells[0].text = str(row["string_label"])
            row_cells[1].text = f"{row['ratio kWh/kWc']:.2f}"

    if inclure_analyse_suspect:
        doc.add_paragraph()
        add_heading2("Analyse des strings suspects")
        if not df_alertes.empty:
            table_alertes = doc.add_table(rows=1, cols=len(df_alertes.columns))
            table_alertes.style = 'Light Grid'
            table_flop.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table_alertes.rows[0].cells
            for i, col in enumerate(df_alertes.columns):
                hdr_cells[i].text = col
            for _, row in df_alertes.iterrows():
                row_cells = table_alertes.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)
        else:
            add_text_paragraph("Aucune alerte détectée.")

    if inclure_evolution:
        doc.add_paragraph()
        add_heading2("Évolution mensuelle")
        add_centered_image(img_evolution)

    temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_path.name)
    return temp_path.name


# Configuration de l'affichage outil
st.set_page_config(
	page_title="String Analyzer", 
	page_icon="🔘",
	layout="wide",
	initial_sidebar_state="expanded"
	)

# Initialisation d'une variable session pour stocker les infos
if "site_name" not in st.session_state:
    st.session_state.site_name = ""
if "fichiers_onduleurs" not in st.session_state:
    st.session_state.fichiers_onduleurs = []
if "fichiers_caracteristiques" not in st.session_state:
    st.session_state.fichiers_caracteristiques = []
if "fichier_irradiance" not in st.session_state:
    st.session_state.fichier_irradiance = None

# Barre latérale pour la navigation
onglet = st.sidebar.radio("", ["💡 Indications","📁 Chargement des données", "📊 Analyse & Visualisation"])

# Corps de l'application


if onglet == "💡 Indications":

    st.title("💡 Indications")


    st.markdown("## Description des fichiers attendus")

    st.markdown("---")

    # ░░░ 1. FICHIER ONDULEUR ░░░
    st.subheader("1️⃣ Fichiers onduleurs")
    st.markdown("Format attendu :")

    
    colonnes_onduleur = ["Time"] + ["String 1"]+["String 2"]+["String 3"] +["..."] +["String N"] + ["total"]
    df_exemple_onduleur = pd.DataFrame([
        ["2025-07-10T08:00:00", 150, 160, 155,'...', 158, '...'],
        ["2025-07-10T08:10:00", 152, 161, 157,'...', 159, '...'],
        ["2025-07-10T08:20:00", 153, 162, 158,'...', 160, '...'],
    ], columns=colonnes_onduleur)

    st.dataframe(df_exemple_onduleur)

    st.markdown("""
    - `time` : date et heure au format `YYYY-MM-DDTHH:MM` (par pas de 10 min)
    - `string x` : puissance du string *x* en **W**
    - `total` : somme des puissances des strings en **W**
    """)

    st.markdown("---")

    # ░░░ 2. FICHIER CARACTÉRISTIQUES ░░░
    st.subheader("2️⃣ Fichiers caractéristiques strings")
    st.markdown("Format attendu :")

    df_exemple_carac = pd.DataFrame([
        {"string": 1, "puissance unitaire": 0.3, "nombre pv": 20},
        {"string": 2, "puissance unitaire": 0.3, "nombre pv": 18},
        {"string": 3, "puissance unitaire": 0.33, "nombre pv": 22},
        {"string": '...', "puissance unitaire": '...', "nombre pv": '...'},
        {"string": 'N', "puissance unitaire": 0.5, "nombre pv": 12},
    ])

    st.dataframe(df_exemple_carac)

    st.markdown("""
    - `string` : numéro du string
    - `puissance unitaire` : puissance d’un panneau du string en **kWc**
    - `nombre pv` : nombre de PV constituant le string
    """)

    st.markdown("---")

    # ░░░ 3. FICHIER IRRADIANCE ░░░
    st.subheader("3️⃣ Fichier irradiance")
    st.markdown("Format attendu :")

    df_exemple_irradiance = pd.DataFrame([
        {"time": "2025-07-10 08:00:00", "irradiance": 0.75},
        {"time": "2025-07-10 08:10:00", "irradiance": 0.77},
        {"time": "2025-07-10 08:20:00", "irradiance": 0.79},
    ])

    st.dataframe(df_exemple_irradiance)

    st.markdown("""
    - `time` : date et heure au format `YYYY-MM-DDTHH:MM' (par pas de 10 min)
    - `irradiance` : irradiance en **kW/m²**
    """)

    st.markdown("---")
    
    
    st.markdown(" ⚠️ Points de vigilance : ")
    st.warning("""
    - 1/ Les fichiers doivent être au **format Excel (.xlsx / .xlsm) ou CSV (.csv)**  
    - 2/ Respecter strictement l’**ordre des colonnes** tel que défini pour chaque tableau 
    - 3/ Veiller à ce que les **strings** soient toujours **ordonnés de manière croissante** ( string 1, string 2, string 3, etc) 
    - 4/ Vérifier que le **format de date** utilisé soit bien conforme à celui requis 
    - 5/ Vérifier que tous les fichiers onduleurs et irradiance couvrent des **périodes** qui coïncident entre elles
    - 6/ S’assurer que toutes les valeurs soient exprimées dans les **unités** demandées
    - 7/ Vérifier que toutes les **valeurs numériques** soient **valides** (aucun caractère non-autorisé ne doit être présent )
    """)    

    st.markdown("A.R.")


elif onglet == "📁 Chargement des données":
    
    st.title("📁 Chargement des données")

    # 1. Entrée du nom du site
    site_name = st.text_input("Nom du site", value=st.session_state.site_name)
    st.session_state.site_name = site_name  # Mémorisation du nom

    # 2. Déclaration du nombre d’onduleurs
    nb_onduleurs = st.number_input("Nombre d’onduleurs", min_value=1, max_value=1000, step=1)

    st.write("")
    st.write("")
    st.write("")


    # 3. Upload des fichiers onduleurs
    st.subheader("🗄️ Données onduleurs")
    fichiers_onduleurs_temp = []
    for i in range(nb_onduleurs):
        fichier = st.file_uploader(f"Onduleur {i+1}", type=["csv", "xlsx"], key=f"onduleur_{i}")
        fichiers_onduleurs_temp.append(fichier)

    if any(f is not None for f in fichiers_onduleurs_temp):
        st.session_state.fichiers_onduleurs = fichiers_onduleurs_temp


    st.write("")
    st.write("")


    # 5. Upload des fichiers caractéristiques strings
    st.subheader("⚡ Données strings")
    fichiers_caracteristiques_temp = []
    for i in range(nb_onduleurs):
        fichier = st.file_uploader(
            f"Caractéristiques strings – Onduleur {i+1}",
            type=["csv", "xlsx"],
            key=f"caracteristiques_{i}"
        )
        fichiers_caracteristiques_temp.append(fichier)
    
    if any(f is not None for f in fichiers_caracteristiques_temp):
        st.session_state.fichiers_caracteristiques = fichiers_caracteristiques_temp
    
    st.write("")
    st.write("")
    
    # 4. Upload du fichier irradiance
    st.subheader("🔆 Données irradiance")
    fichier_irradiance_temp = st.file_uploader("",type=["csv", "xlsx"], key="irradiance")

    if fichier_irradiance_temp is not None:
        st.session_state.fichier_irradiance = fichier_irradiance_temp


    st.write("")
    st.write("")
    st.write("")


    # 6. Bouton de validation
    if st.button("Valider"):
        
        # Vérification des fichiers onduleurs
        if any(f is None for f in st.session_state.fichiers_onduleurs):
            st.warning("⚠️ Tous les fichiers d’onduleurs n’ont pas été importés.")

        # Vérification du fichier irradiance
        elif st.session_state.fichier_irradiance is None:
            st.error("⚠️ Le fichier d’irradiance n’a pas été importé.")
            
        # Vérification des fichier caractéristiques strings
        elif any(f is None for f in st.session_state.fichiers_caracteristiques):
            st.error("⚠️ Le fichier des caractéristiques des strings n’a pas été importé.")

        else:
            # Vérification cohérence colonnes ↔ lignes
            coherence_valide = True  
            for i in range(nb_onduleurs):
                try:
                    df_onduleur = lire_fichier(st.session_state.fichiers_onduleurs[i])
                    df_carac = lire_fichier(st.session_state.fichiers_caracteristiques[i])

                    nb_colonnes_utiles = len(df_onduleur.columns) - 2
                    nb_lignes_carac = len(df_carac)

                    if nb_colonnes_utiles != nb_lignes_carac:
                        st.error(
                            f" ❌  Incohérence pour l'onduleur {i+1} :  "
                            f"{nb_colonnes_utiles} strings détectés dans le fichier onduleur, "
                            f"mais {nb_lignes_carac} strings dans le fichier de caractéristiques."
                        )
                        coherence_valide = False

                except Exception as e:
                    st.error(f"Erreur lors de la lecture des fichiers pour l'onduleur {i+1} : {e}")
                    coherence_valide = False

        # Validation des fichiers uploadés si toutes les conditions sont respectées
        if coherence_valide:
            st.success(f"Données du site {site_name} chargées avec succès")


elif onglet == "📊 Analyse & Visualisation":

    # Vérification de la présence des fichiers requis
    if (any(f is None for f in st.session_state.fichiers_onduleurs)
    or st.session_state.fichier_irradiance is None
    or any(f is None for f in st.session_state.fichiers_caracteristiques)):
        
        st.error("Importer tous les fichiers nécessaires pour accéder à cet onglet.")
    else:

        st.title(f"📊 Analyse & Visualisation {st.session_state.site_name}")


        # Menu déroulant pour "Etude globale"
        with st.expander("ETUDE GLOBALE", expanded=False):

            # 1. Choix de l'onduleur à analyser
            onduleur_index = st.selectbox("Choisir un onduleur à analyser", range(1, len(st.session_state.fichiers_onduleurs)+1), format_func=lambda x: f"Onduleur {x}")
            
            # 2. Chargement des données correspondant à cet onduleur
            fichier_onduleur = st.session_state.fichiers_onduleurs[onduleur_index-1]
            fichier_carac = st.session_state.fichiers_caracteristiques[onduleur_index-1]
            fichier_irradiance = st.session_state.fichier_irradiance

            if fichier_onduleur and fichier_carac and fichier_irradiance:

                # 3. Lecture et nettoyage des fichiers

                    #  onduleur
                df_puissance = traiter_fichier_onduleur(fichier_onduleur)

                    # Données strings
                df_carac = traiter_fichier_carac(fichier_carac)
                
                    # Données irradiance
                df_irradiance = traiter_fichier_irradiance(fichier_irradiance)
            
                # 4. Choix de la période d’analyse
                min_date = max(df_puissance["time"].min().date(), df_irradiance["time"].min().date())
                max_date = min(df_puissance["time"].max().date(), df_irradiance["time"].max().date())

                col1, col2 = st.columns(2)
                with col1:
                    date_debut = st.date_input("📅 Date de début", min_value=min_date, max_value=max_date, value=min_date)
                with col2:
                    date_fin = st.date_input("📅 Date de fin", min_value=min_date, max_value=max_date, value=max_date)
                
                if date_fin < date_debut:
                    st.error("La date de fin doit être supérieure ou égale à la date de début.")
                    st.stop()

                # 5. Filtrage des données par période
                df_puissance_filtré = df_puissance[(df_puissance["time"].dt.date >= date_debut) & (df_puissance["time"].dt.date <= date_fin)]
                df_irradiance_filtré = df_irradiance[(df_irradiance["time"].dt.date >= date_debut) & (df_irradiance["time"].dt.date <= date_fin)]

                
                # 6. Calcul des données utiles
                    
                    # Calcul des puissances réelles
                colonnes_strings = [col for col in df_puissance_filtré.columns if col not in ["time"]]
                puissances_reelles= df_puissance_filtré[colonnes_strings]/1000
                puissances_reelles_moyenne = puissances_reelles.mean(axis=0) 

                    # Calcul des puissances théoriques
                df_merged = pd.merge_asof(
                df_puissance_filtré[["time"]],
                df_irradiance_filtré,
                on="time",
                direction="nearest"
                )

                puissances_theoriques = df_merged[["time"]].copy()
                    
                for _, row in df_carac.iterrows():
                    string_name = f"string {int(row['string'])}"
                    unitaire = float(row["puissance unitaire"])
                    n_pv = int(row["nombre pv"])
                    puissances_theoriques[string_name] = df_merged["irradiance"] * n_pv * unitaire * 0.8
                        
                puissances_theoriques_moyenne = puissances_theoriques.drop(columns="time").mean(axis=0)
                puissances_theoriques_moyenne["total"] = puissances_theoriques_moyenne.sum()
                
                    # Calcul des énergies réelles et théoriques
                energies_reelles = puissances_reelles.sum(axis=0) * (10/60)
                energies_theoriques = puissances_theoriques.drop(columns="time").sum(axis=0) * (10/60)
                energies_theoriques["total"] = energies_theoriques.sum()

                
                # 7. Alignement des deux séries
                index_communs = sorted(set(puissances_reelles_moyenne.index) & set(puissances_theoriques_moyenne.index))
                puissances_reelles_moyenne = puissances_reelles_moyenne.reindex(index_communs)
                puissances_theoriques_moyenne = puissances_theoriques_moyenne.reindex(index_communs)

                    
                
                option_etude = st.radio(" Choisir l’analyse à afficher :", ["🔍 Données générales","🔍 Puissance moyenne réelle vs théorique (kW)","🔍 Énergie totale réelle vs théorique (kWh)"], horizontal=False)
                
                # 8. Afficher l'analyse sur les données générales
                if option_etude == "🔍 Données générales":

                    # Affichage du tableau de puissance  
                    st.subheader(" Puissances (kW)")
                    df_puissance_affiche = df_puissance_filtré[["time"]].copy()
                    df_puissance_affiche[colonnes_strings] = puissances_reelles
                    st.dataframe(df_puissance_affiche.style.format({col: "{:.2f}" for col in colonnes_strings}))
                    

                    # Affichage du tableau d'énergie  
                    st.subheader(" Énergies (kWh) ")
                    df_energie_affiche = df_puissance_filtré[["time"]].copy() 
                    df_energie_affiche[colonnes_strings] = puissances_reelles[colonnes_strings] * (10/60)  # kWh = kW × h
                    st.dataframe(df_energie_affiche.style.format({col: "{:.2f}" for col in colonnes_strings}))
                
                # 9. Afficher l'analyse sur la puissance moyenne réelle vs théorique
                elif option_etude == "🔍 Puissance moyenne réelle vs théorique (kW)":
                        
                    # Création d'un dataframe de puissance
                    df_puissance_chart = pd.DataFrame({
                        "P. moyenne réelle (kW)": puissances_reelles_moyenne,
                        "P. moyenne théorique (kW)": puissances_theoriques_moyenne,
                        }).reset_index().rename(columns={"index": "string"})

                    df_puiss_long = df_puissance_chart.melt(id_vars="string", var_name="type", value_name="valeur")
                    categories_triees = trier_strings_par_numero([str(s) for s in df_puiss_long["string"].unique()])
                    df_puiss_long["string"] = pd.Categorical(df_puiss_long["string"], categories=categories_triees, ordered=True)

                    # Affichage du graphique de puissance
                    graph_width = max(700, len(df_puissance_chart) * 50)
                    graph_height = max(400, len(df_puissance_chart) * 25)

                    chart_puissance = alt.Chart(df_puiss_long).mark_bar().encode(
                    x=alt.X('string:N',sort=categories_triees,title="String"),
                    xOffset='type:N',
                    y=alt.Y('valeur:Q', title="Puissance (kW)"),
                    color=alt.Color('type:N',scale=alt.Scale(domain=["P. moyenne réelle (kW)", "P. moyenne théorique (kW)"],range=["#1f77b4", "#ff7f0e"]),title="Légende"),
                    tooltip=['string', 'type', 'valeur']).properties(width=graph_width,height=graph_height,title="")

                    st.altair_chart(chart_puissance, use_container_width=True)

                    # Affichage du tableau recap
                    with st.expander("📋 Détails",expanded=False):
                        st.dataframe(df_puissance_chart.style.format({col: "{:.2f}" for col in df_puissance_chart.select_dtypes(include="number").columns}))

                
                # 10. Afficher l'analyse sur l'énergie totale réelle vs théorique
                elif option_etude =="🔍 Énergie totale réelle vs théorique (kWh)":

                    # Réindexation des données
                    energies_reelles = energies_reelles.reindex(index_communs)
                    energies_theoriques = energies_theoriques.reindex(index_communs)    

                    # Création d'un dataframe d'énergie  
                    df_energie_chart = pd.DataFrame({
                    "E. totale réelle (kWh)": energies_reelles,
                    "E. totale théorique (kWh)": energies_theoriques
                    }).reset_index().rename(columns={"index": "string"})

                    df_energie_long = df_energie_chart.melt(id_vars="string", var_name="type", value_name="valeur")
                    categories_triees = trier_strings_par_numero([str(s) for s in df_energie_long["string"].unique()])
                    df_energie_long["string"] = pd.Categorical(df_energie_long["string"], categories=categories_triees, ordered=True)
                    
                    graph_width = max(700, len(df_energie_chart) * 50)
                    graph_height = max(400, len(df_energie_chart) * 25)
                    
                    # Affichage du graphique d'énergie
                    chart_energie = alt.Chart(df_energie_long).mark_bar().encode(
                        x=alt.X('string:N',sort=categories_triees, title="String"),
                        xOffset='type:N',
                        y=alt.Y('valeur:Q', title="Energie (kWh)"),
                        color=alt.Color('type:N',scale=alt.Scale(domain=["E. totale réelle (kWh)", "E. totale théorique (kWh)"],range=["#2ca02c", "#d62728"]),title="Type de donnée"),
                        tooltip=['string', 'type', 'valeur']).properties(width=graph_width, height=graph_height)

                    st.altair_chart(chart_energie, use_container_width=True)

                    # Affichage du tableau recap
                    with st.expander("📋 Détails",expanded=False):
                        st.dataframe(df_energie_chart.style.format({col: "{:.2f}" for col in df_energie_chart.select_dtypes(include="number").columns}))
                            
                else:
                    st.warning("Choisir l'analyse à afficher")           

            else:
                st.warning("Merci de charger tous les fichiers.")
        
        # Menu déroulant pour "Suivi temporel de la puissance"
        with st.expander("SUIVI TEMPOREL DE LA PUISSANCE", expanded=False):

            # 1. Choix de l'onduleur à analyser
            onduleur_index = st.selectbox("Choisir un onduleur", range(1, len(st.session_state.fichiers_onduleurs)+1), format_func=lambda x: f"Onduleur {x}", key="suivi_onduleur_index")

            # 2. Chargement des données correspondant à cet onduleur
            fichier_onduleur = st.session_state.fichiers_onduleurs[onduleur_index-1]

            st.write("")

            if fichier_onduleur:

                # 3. Lecture et nettoyage des données
                df_puissance = traiter_fichier_onduleur(fichier_onduleur)

                # 4. Détection plage de dates
                min_date = df_puissance["time"].min().date()
                max_date = df_puissance["time"].max().date()

                date_choisie = st.date_input("📅 Choisir un jour", min_value=min_date, max_value=max_date, value=min_date, key="jour_analyse")

                # 5. Filtrage pour le jour sélectionné
                df_jour = df_puissance[df_puissance["time"].dt.date == date_choisie]

                st.write("")

                if not df_jour.empty:

                    # 6. Configuration des options d'affichage
                    strings_disponibles = [col for col in df_jour.columns if col not in ["time"]]
                    options = ["Tout"] + strings_disponibles
                    sélection = st.multiselect("Sélectionner les strings à afficher :", options=options, default=["Tout"], key="multiselect_strings_evolution")

                    if "Tout" in sélection:
                        strings_affichées = strings_disponibles
                    else:
                        strings_affichées = sélection

                    # 7. Mise en forme pour l'afichage graphique
                    df_plot = df_jour[["time"] + strings_affichées].melt(id_vars="time", var_name="string", value_name="puissance")
                    categories_triees = trier_strings_par_numero([str(s) for s in df_plot["string"].unique()])
                    df_plot["string"] = pd.Categorical(df_plot["string"], categories=categories_triees, ordered=True)

                

                    # 8. Affichage du graphique d'évolution temporelle de la puissance
                    chart = alt.Chart(df_plot).mark_line().encode(
                        x=alt.X("time:T",sort=categories_triees, title="Temps",axis=alt.Axis(format="%H:%M",tickMinStep=3600000)),
                        y=alt.Y("puissance:Q", title="Puissance (kW)"),
                        color=alt.Color("string:N", title="String",sort=categories_triees),
                        tooltip=["time", "string", "puissance"]
                    ).properties(
                        width=900,
                        height=400,
                        title=f"Puissance au cours du temps – {date_choisie.strftime('%d/%m/%Y')}"
                    ).interactive()

                    st.write("")
                    st.write("")

                    st.altair_chart(chart, use_container_width=True)
                else:
                    st.warning("Aucune donnée disponible pour ce jour.")

        # Menu déroulant pour "Performance"
        with st.expander("PERFORMANCE", expanded=False):

            
            mode_perf = st.radio("**Choisir l’analyse à effectuer**", ["🔍 Étude par onduleur","🔍 Comparaison entre onduleurs"],horizontal=False)
            
            st.write("")

            # 1. Affichage de l'analyse sur l'étude par onduleur            
            if mode_perf == "🔍 Étude par onduleur":

                # Sélection de l'onduleur à analyser
                onduleur_index = st.selectbox("Choisir un onduleur", range(1, len(st.session_state.fichiers_onduleurs)+1), format_func=lambda x: f"Onduleur {x}", key="perf_onduleur_index")
                fichier_onduleur = st.session_state.fichiers_onduleurs[onduleur_index-1]
                fichier_carac = st.session_state.fichiers_caracteristiques[onduleur_index-1]

                if fichier_onduleur and fichier_carac:

                    df_puissance = traiter_fichier_onduleur(fichier_onduleur)
                    
                    colonnes_strings = [col for col in df_puissance.columns if (col != "time") and (col != "total")]
                    df_puissance[colonnes_strings] = df_puissance[colonnes_strings] / 1000  # W → kW

                    # Sélection de la période à analyser
                    min_date = df_puissance["time"].min().date()
                    max_date = df_puissance["time"].max().date()

                    col1, col2 = st.columns(2)
                    with col1:
                        date_debut = st.date_input("📅 Date de début", min_value=min_date, max_value=max_date, value=min_date, key="perf_start")
                    with col2:
                        date_fin = st.date_input("📅 Date de fin", min_value=min_date, max_value=max_date, value=max_date, key="perf_end")

                    st.write(f" Données disponibles de {min_date} à {max_date}")
                    
                    if date_fin < date_debut:
                        st.error("La date de fin doit être supérieure ou égale à la date de début.")
                        st.stop()
                
                    # Filtrage des données
                    df_filtré = df_puissance[(df_puissance["time"].dt.date >= date_debut) & (df_puissance["time"].dt.date <= date_fin)]
                    
                    st.write("")
                    st.write("")
                    
                    if not df_filtré.empty:

                        # Lecture et nettoyage des données
                        df_carac = traiter_fichier_carac(fichier_carac)
                        df_carac["puissance installée (kWc)"] = df_carac["puissance unitaire"] * df_carac["nombre pv"] 

                        # Calcul de l'énergie réelle par string (en kWh)
                        df_energie = df_filtré.copy()
                        df_energie[colonnes_strings] = df_energie[colonnes_strings] * (10/60)
                        energie_totale = df_energie[colonnes_strings].sum(axis=0)

                        
                        # Création du dataframe d'énergie
                        df_resultats = pd.DataFrame({
                        "string": [int(s.split()[-1]) for s in energie_totale.index],
                        "energie produite (kWh)": energie_totale.values})

                        # Jointure avec les caractéristiques
                        df_resultats = df_resultats.merge(df_carac[["string", "puissance installée (kWc)"]], on="string", how="left")

                        # Calcul des ratios de performance
                        df_resultats["ratio kWh/kWc"] = df_resultats["energie produite (kWh)"] / df_resultats["puissance installée (kWc)"]

                        # Affichage du graphe des ratios de performance
                        categories_triees = trier_strings_par_numero([f"string {i}" for i in df_resultats["string"]])
                        df_resultats["string_label"] = [f"string {i}" for i in df_resultats["string"]]
                        df_resultats["string_label"] = pd.Categorical(df_resultats["string_label"], categories=categories_triees, ordered=True)

                        st.subheader("📶 Ratio kWh / kWc par string")

                        graph_width = max(700, len(df_resultats) * 50)
                        graph_height = max(400, len(df_resultats) * 25)

                        chart_ratio = alt.Chart(df_resultats).mark_bar().encode(
                            x=alt.X("string_label:N", sort=categories_triees, title="String"),
                            y=alt.Y("ratio kWh/kWc:Q", title="Performance (kWh / kWc)"),
                            tooltip=["string_label", "ratio kWh/kWc"]).properties(width=graph_width, height=graph_height)

                        st.altair_chart(chart_ratio, use_container_width=True)

                        with st.expander("📋 Détails",expanded=False):
                            df_récap = df_resultats[["string_label", "ratio kWh/kWc"]].rename(columns={"string_label": "string"}).sort_values("string")
                            st.dataframe(df_récap.style.format({"ratio kWh/kWc": "{:.2f}"}), use_container_width=True)
                        
                        # Tri des données
                        df_trie_desc = df_resultats.sort_values("ratio kWh/kWc", ascending=False).reset_index(drop=True)
                        df_trie_asc = df_resultats.sort_values("ratio kWh/kWc", ascending=True).reset_index(drop=True)

                        # Création des tableaux de top 3 et bottom 3
                        n_top = min(3, len(df_trie_desc))
                        n_flop = min(3, len(df_trie_asc))

                        top3 = df_trie_desc.head(n_top)
                        flop3 = df_trie_asc.head(n_flop)

                        st.write("")
                        st.write("")

                        # Affichage des tableaux sur les classement des strings obtenus
                        st.subheader("🔻 Classement des strings")
                        col1, col2 = st.columns(2)
                        with col1:
                            st.markdown("**Top 3 - Strings plus performants**")
                            st.dataframe(top3[["string_label","ratio kWh/kWc"]].rename(columns={"string_label": "string"}).reset_index(drop=True).style.format({"ratio kWh/kWc":"{:.2f}"}))

                        with col2:
                            st.markdown("**Bottom 3 - Strings moins performants**")
                            st.dataframe(flop3[["string_label","ratio kWh/kWc"]].rename(columns={"string_label": "string"}).reset_index(drop=True).style.format({"ratio kWh/kWc":"{:.2f}"}))
        
                        st.write("")
                        st.write("")


                        
                        k=1.5
                        alertes = []

                        # Calcul moyenne globale des ratios
                        moyenne_globale = df_resultats["ratio kWh/kWc"].mean()

                        # Sélection des strings sous la moyenne
                        strings_sous_moyenne = df_resultats[df_resultats["ratio kWh/kWc"] < moyenne_globale]
                        print(df_resultats)

                        for idx, row in strings_sous_moyenne.iterrows():
                            string_test = row["string"]
                            ratio_test = row["ratio kWh/kWc"]

                            # Récupération des caractéristiques du string
                            carac_test = df_carac[df_carac["string"] == string_test].iloc[0]
                            nb_pv_test = carac_test["nombre pv"]
                            pu_test = carac_test["puissance unitaire"]

                            # Recherche des strings comparables
                            strings_comparables = df_carac[
                                (df_carac["string"] != string_test) &
                                (df_carac["nombre pv"] == nb_pv_test) &
                                (df_carac["puissance unitaire"] == pu_test)
                            ]["string"].tolist()

                            # Récupération des ratios des strings comparables

                            ratios_comparables = df_resultats[df_resultats["string_label"].isin([f"string {i}" for i in strings_comparables])]["ratio kWh/kWc"] if strings_comparables else pd.Series()
                            
                            # Mis en place du message adequat
                            if not ratios_comparables.empty:
                                moyenne = ratios_comparables.mean()
                                ecart_type = ratios_comparables.std()
                                seuil_alerte = moyenne - (k * ecart_type)
                                
                                message = (
                                    "🔴 Anormal"
                                    if ratio_test < seuil_alerte
                                    else "🟡 Acceptable"
                                )
                                ecart_pct = np.abs(moyenne-ratio_test)/moyenne * 100
                            else:
                                message = "Pas d’éléments de comparaison"
                                ecart_pct = "—"

                            alertes.append({
                                "String": f"string {string_test}",
                                "Écart à la moyenne (%)": f"{ecart_pct:.2f}",
                                "Message": message
                            })

                        

                        # Affichage des alertes
                        if alertes:
                            st.subheader("🚨 Analyse des strings suspects")
                            st.caption("(Ratio inférieur à la moyenne globale)")
                            df_alertes = pd.DataFrame(alertes)
                            st.dataframe(df_alertes.style.applymap(
                                lambda val: 'color: red; font-weight: bold' if isinstance(val, str) and "Alerte" in val else '',
                                subset=["Message"]
                            ))
                        
                        st.write("")
                        st.write("")

                        st.subheader("📉 Evolution mensuelle ")

                        # Filtrage des données sur la période choisie
                        df_puissance = df_puissance[(df_puissance["time"].dt.date >= date_debut) & (df_puissance["time"].dt.date <= date_fin)]

                        # Ajout d'une colonne pour regrouper les données par mois
                        df_puissance["year_month"] = df_puissance["time"].dt.to_period("M")

                        # Calcul de l'énergie par string par mois 
                        df_energie = df_puissance.copy()
                        df_energie[colonnes_strings] = df_energie[colonnes_strings] * (10 / 60)  # puissance * durée en h

                        # Calcul de la somme de l'énergie par string par mois
                        df_mensuel = df_energie.groupby("year_month")[colonnes_strings].sum().reset_index()
                
                        # Calcul ratio mensuel de performance
                        ratios_mensuels = []
                        for s in colonnes_strings:
                            num_string = int(s.split()[-1])
                            puissance_string = df_carac.loc[df_carac["string"] == num_string, "puissance installée (kWc)"].values
                            if len(puissance_string) == 0:
                                continue
                            puissance_string = puissance_string[0]
                            df_tmp = df_mensuel[["year_month", s]].copy()
                            df_tmp.rename(columns={s: "energie produite (kWh)"}, inplace=True)
                            df_tmp["ratio kWh/kWc"] = df_tmp["energie produite (kWh)"] / puissance_string
                            df_tmp["string"] = s
                            ratios_mensuels.append(df_tmp)

                        df_ratios = pd.concat(ratios_mensuels, ignore_index=True)

                        # Conversion des données de year_month en datetime 
                        df_ratios["year_month"] = df_ratios["year_month"].dt.to_timestamp()

                        # Tri des strings disponibles selon leur nom
                        strings_disponibles = trier_strings_par_numero([str(s) for s in df_ratios["string"].unique()])
                        
                        
                        # Configuration des options d'affichage  
                        options = ["Tout"] + strings_disponibles
                        sélection = st.multiselect(
                        "Sélectionner les strings à afficher :",
                        options=options,
                        default=["Tout"],
                        key="multiselect_strings_degradation")

                    
                        if "Tout" in sélection or not sélection:
                            strings_affichées = strings_disponibles
                        else:
                            strings_affichées = sélection

                        # Filtrage du dataframe
                        df_affichage = df_ratios[df_ratios["string"].isin(strings_affichées)]


                        # Affichage du graphique d'évolution temporelle des ratios de performance
                        chart = alt.Chart(df_affichage).mark_line(point=True).encode(
                            x=alt.X("year_month:T",sort=strings_disponibles, title="Mois",axis=alt.Axis(format='%b',tickCount="month")),
                            y=alt.Y("ratio kWh/kWc:Q", title=" Ratio kWh / kWc"),
                            color=alt.Color("string:N", title="String", sort=strings_disponibles),
                            tooltip=["string", alt.Tooltip("year_month:T", title="Mois",format="%b %Y"), alt.Tooltip("ratio kWh/kWc")]
                        ).properties(width=800, height=400)

                        st.altair_chart(chart, use_container_width=True)

                        with st.expander("📋 Details", expanded=False):

                            # Création du tableau récap
                            tableau_croisé = df_affichage.pivot_table(
                                index=df_affichage["year_month"].dt.strftime("%Y-%m"),  
                                columns="string",  
                                values="ratio kWh/kWc",  
                                aggfunc="mean" 
                            )
                            
                            tableau_croisé.index.name = "date" 
                            tableau_croisé = tableau_croisé.sort_index()  

                            colonnes_ordre = trier_strings_par_numero([str(c) for c in tableau_croisé.columns])
                            tableau_croisé = tableau_croisé[colonnes_ordre]
                            tableau_croisé_formaté = tableau_croisé.style.format("{:.2f}")

                            st.dataframe(tableau_croisé_formaté, use_container_width=True)

                        st.write("")
                        st.write("")

                # Bouton de génération
        
                st.markdown("📑 **Sections à inclure dans le rapport**")
                inclure_ratio = st.checkbox("Ratio kWh/kWc par string", value=True)
                inclure_classement = st.checkbox("Classement des strings (Top/Bottom 3)", value=True)
                inclure_analyse_suspect = st.checkbox("Analyse des strings suspects", value=True)
                inclure_evolution = st.checkbox("Évolution mensuelle", value=True)
                    
                if st.button ("Générer le rapport") :
                    site = st.session_state.site_name
                    onduleur = onduleur_index
                    debut = date_debut.strftime("%Y-%m-%d")
                    fin = date_fin.strftime("%Y-%m-%d")

                    # Fichiers temporaires pour images PNG
                    img_barres = tempfile.NamedTemporaryFile(delete=False, suffix=".png").name
                    img_evolution = tempfile.NamedTemporaryFile(delete=False, suffix=".png").name

                    # Sauvegarde des graphiques avec CairoSVG
                    sauvegarder_chart_png(chart_ratio, img_barres)
                    sauvegarder_chart_png(chart, img_evolution)

                    # Génération du rapport 
                    chemin_rapport = generer_word(
                        site, onduleur, debut, fin,
                        img_barres, top3, flop3, pd.DataFrame(alertes), img_evolution,
                        inclure_ratio=inclure_ratio,
                        inclure_classement=inclure_classement,
                        inclure_analyse_suspect=inclure_analyse_suspect,
                        inclure_evolution=inclure_evolution,
                        logo_path="logo_NEA.png"
                    )
                    
                    # Proposition de téléchargement 
                    with open(chemin_rapport, "rb") as f:
                        st.download_button(
                            label=f"📥 Télécharger le rapport",
                            data=f,
                            file_name=f"Rapport_{site}_performance_onduleur{onduleur}_{debut}_{fin}.docx"
                        )

            # 2. Affichage de l'analyse sur la comparaison entre onduleurs    
            elif mode_perf == "🔍 Comparaison entre onduleurs":

                # Récupération de la période commune à tous les onduleurs
                dates_min = []
                dates_max = []
                for f_onduleur in st.session_state.fichiers_onduleurs:
                    if f_onduleur is None:
                        continue
                    df_temp = traiter_fichier_onduleur(f_onduleur)
                    dates_min.append(df_temp["time"].min().date())
                    dates_max.append(df_temp["time"].max().date())
                
                if not dates_min or not dates_max:
                    st.warning("Impossible de déterminer la période commune.")
                    st.stop()
                
                date_debut_min = max(dates_min)  
                date_fin_max = min(dates_max)    

                # Choix de la période commune
                col1, col2 = st.columns(2)
                with col1:
                    date_debut = st.date_input("📅 Date de début", min_value=date_debut_min, max_value=date_fin_max, value=date_debut_min, key="comp_date_debut")
                with col2:
                    date_fin = st.date_input("📅 Date de fin", min_value=date_debut_min, max_value=date_fin_max, value=date_fin_max, key="comp_date_fin")

                if date_fin < date_debut:
                    st.error("La date de fin doit être postérieure ou égale à la date de début.")
                    st.stop()
                
                st.write("")
                st.write("")
                    

                # Calcul des ratios filtrés par période pour chaque onduleur
                all_ratios = []
                    
                for idx_onduleur, (f_onduleur, f_carac) in enumerate(zip(st.session_state.fichiers_onduleurs, st.session_state.fichiers_caracteristiques), start=1):
                    if f_onduleur is None or f_carac is None:
                        st.warning(f"Fichiers manquants pour l'onduleur {idx_onduleur}")
                        continue
                    
                    # Lecture et nettoyage des fichiers onduleurs
                    df_puissance = traiter_fichier_onduleur(f_onduleur)
                

                    # Filtrage sur la période choisie
                    df_puissance = df_puissance[(df_puissance["time"].dt.date >= date_debut) & (df_puissance["time"].dt.date <= date_fin)]

                    # Sélection des colonnes de strings 
                    colonnes_strings = [col for col in df_puissance.columns if col not in ["time", "total"]]
                    df_puissance[colonnes_strings] = df_puissance[colonnes_strings] / 1000  

                    # Calcul de la puissance installée pour chaque string
                    df_carac = traiter_fichier_carac(f_carac)
                    df_carac["puissance installée (kWc)"] = df_carac["puissance unitaire"] * df_carac["nombre pv"]

                    # Calcul de l’énergie produite pour chaque string
                    df_energie = df_puissance.copy()
                    df_energie[colonnes_strings] = df_energie[colonnes_strings] * (10 / 60) 
                    energie_totale = df_energie[colonnes_strings].sum(axis=0)

                    df_resultats = pd.DataFrame({
                        "string": [int(s.split()[-1]) for s in energie_totale.index],
                        "energie produite (kWh)": energie_totale.values
                    })

                    # Calcul des ratios de performance
                    df_resultats = df_resultats.merge(df_carac[["string", "puissance installée (kWc)"]], on="string", how="left")
                    df_resultats["ratio kWh/kWc"] = df_resultats["energie produite (kWh)"] / df_resultats["puissance installée (kWc)"]

                    df_resultats["onduleur"] = f"Onduleur {idx_onduleur}"
                    all_ratios.append(df_resultats[["string", "ratio kWh/kWc", "onduleur"]])
                
                # Affichage du graphique de ratios de performance
                if all_ratios:
                    df_comparaison = pd.concat(all_ratios, ignore_index=True)

                    df_comparaison["string_label"] = "string " + df_comparaison["string"].astype(str)
                    categories_triees = trier_strings_par_numero(df_comparaison["string_label"].unique())
                    df_comparaison["string_label"] = pd.Categorical(df_comparaison["string_label"], categories=categories_triees, ordered=True)

                    
                    graph_width = max(700, len(categories_triees) * 50)
                    graph_height = max(400, len(categories_triees) * 25)
                    

                    chart = alt.Chart(df_comparaison).mark_bar().encode(
                        x=alt.X("string_label:N", sort=categories_triees, title="String"),
                        y=alt.Y("ratio kWh/kWc:Q", title="Performance (kWh / kWc)"),
                        color=alt.Color("onduleur:N", title="Onduleur"),
                        xOffset="onduleur:N",
                        tooltip=["onduleur", "string_label", alt.Tooltip("ratio kWh/kWc")]).properties(width=graph_width, height=graph_height, title="Ratios kWh/kWc par string et onduleur")

                    st.altair_chart(chart, use_container_width=True)

                    with st.expander("📋Détails", expanded=False):
                        tableau_croisé = df_comparaison.pivot_table(
                        index="string_label",
                        columns="onduleur",
                        values="ratio kWh/kWc",
                        aggfunc="mean"
                    )
                        tableau_croisé.index.name = "string"  
                        tableau_croisé = tableau_croisé.round(2).sort_index()
                        tableau_croisé_formaté = tableau_croisé.style.format("{:.2f}")

                        st.dataframe(tableau_croisé_formaté, use_container_width=True)
                        
                else:
                    st.warning("Aucune donnée complète pour comparaison entre onduleurs.")




# Code écrit par Amboara RASOLOFOARIMANANA